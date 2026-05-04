from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "StudySync/docs/client/StudySync_Client_Architecture_Report_정태현.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = "Body Text"
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        rest = text[len(bold_prefix):]
        p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10)
    return p


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F6FA")
    p = cell.paragraphs[0]
    for line in text.strip("\n").splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "1F4E79")
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(5)

    for idx, size, color in [(1, 16, "1F4E79"), (2, 13, "2F75B5"), (3, 11, "385723")]:
        style = styles[f"Heading {idx}"]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build():
    doc = Document()
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("StudySync Client Architecture Report")
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("실시간 머신비전 비동기 파이프라인 설계 및 진행 현황")
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(68, 68, 68)

    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["작성자", "정태현"],
            ["대상 모듈", "StudySync / client"],
            ["작성일", "2026-05-04"],
            ["목적", "팀원에게 클라이언트 설계 의도, 파이프라인 구조, 현재 구현 진행도를 공유"],
            ["현재 단계", "MFC/C++ 클라이언트 아키텍처 뼈대 구현 및 통신 책임 분리 완료"],
        ],
        widths=[3.0, 12.5],
    )

    add_heading(doc, "1. 설계 목표", 1)
    add_body(doc, "StudySync 클라이언트는 웹캠 기반 실시간 자세 분석 앱으로, 화면 렌더링과 AI 분석, 로그 전송, 이벤트 클립 저장을 서로 분리하여 네트워크 지연이 화면 출력에 영향을 주지 않도록 설계한다.")
    add_body(doc, "핵심 목표는 30fps 이상의 부드러운 로컬 렌더링, 5fps 샘플링 기반 AI 서버 분석, 이벤트 발생 시점의 클립 추출, JSONL 기반 경량 로그 전송이다.")

    add_heading(doc, "2. 전체 파이프라인", 1)
    add_code_block(
        doc,
        """
CaptureThread
  -> render_buffer_ : RingBuffer<Frame, 8>        실시간 화면 렌더링
  -> send_buffer_   : RingBuffer<Frame, 8>        AI 서버 5fps 샘플링 전송
  -> shadow_buffer_ : EventShadowBuffer<Frame,60> 이벤트 클립 추출용 최근 프레임 보관

RenderThread
  -> wait_pop(render_buffer_)
  -> Direct2D bitmap upload
  -> OverlayPainter

ZmqSendThread
  -> try_pop(send_buffer_)
  -> 6프레임마다 1장 JPEG 인코딩
  -> IFrameSender(ZmqFrameSender)

ZmqRecvThread
  -> pose_server.py 분석 결과 수신
  -> PostureEventDetector
  -> EventQueue

EventUploadThread
  -> IEventClipStore(LocalClaimCheckClipStore)
  -> ILogSink(HttpJsonlLogSink)
        """,
    )

    add_heading(doc, "3. 핵심 설계 결정", 1)
    add_table(
        doc,
        ["영역", "설계", "이유"],
        [
            ["렌더링", "Direct2D 기반 2D 렌더링", "MFC/GDI보다 영상 프레임과 오버레이를 안정적으로 출력하기 쉬움"],
            ["프레임 흐름", "고정 크기 RingBuffer", "오래된 프레임을 덮어써 실시간성을 유지하고 지연 누적을 방지"],
            ["AI 분석", "ZMQ 기반 5fps 샘플링", "분석 서버 지연이 렌더링을 막지 않도록 분리"],
            ["로그", "HTTP JSONL", "분석 결과와 이벤트 메타데이터를 줄 단위로 배치 전송"],
            ["영상 클립", "Claim Check Pattern", "이벤트 영상 원본은 큐에 직접 넣지 않고 저장 경로만 로그로 전송"],
            ["확장성", "인터페이스 기반 통신 주입", "ZMQ/HTTP/저장 방식을 함수 인자와 설정값으로 교체 가능"],
        ],
        widths=[2.6, 4.5, 8.4],
    )

    add_heading(doc, "4. 통신 책임 분리", 1)
    add_body(doc, "통신은 하나의 거대한 Sender로 묶지 않고 목적별 인터페이스로 분리했다. 이 방식은 단일 책임 원칙을 지키며, 추후 통신 방식 변경 시 핵심 파이프라인을 수정하지 않고 구현체만 교체할 수 있게 한다.")
    add_code_block(
        doc,
        """
IFrameSender
  - Frame 전송 책임
  - 현재 구현체: ZmqFrameSender

ILogSink
  - JSONL 로그와 이벤트 메타데이터 전송 책임
  - 현재 구현체: HttpJsonlLogSink

IEventClipStore
  - 이벤트 클립 저장 및 참조 URI 반환 책임
  - 현재 구현체: LocalClaimCheckClipStore

ClientTransportConfig
  - ZMQ endpoint, JSONL endpoint, clip directory, sample interval, JPEG quality 설정
        """,
    )

    add_heading(doc, "5. 현재 구현 진행도", 1)
    add_table(
        doc,
        ["구분", "상태", "진행률", "설명"],
        [
            ["프로젝트 구조", "완료", "90%", "MFC App/MainFrame/View, Visual Studio 프로젝트, CMake 진입점 구성"],
            ["프레임 버퍼", "완료", "85%", "RingBuffer, ThreadSafeQueue, render/send/shadow buffer 분리"],
            ["캡처 스레드", "초안 구현", "70%", "OpenCV VideoCapture 기반 캡처 루프와 3개 버퍼 push 구조 구현"],
            ["렌더 스레드", "초안 구현", "55%", "wait_pop 기반 렌더 스레드 구조 구현, Direct2D 업로더 세부 구현 필요"],
            ["AI 전송", "뼈대 구현", "55%", "ZmqSendThread, IFrameSender, ZmqFrameSender 구조 구현. 실제 ZMQ socket 송신 TODO"],
            ["AI 수신", "뼈대 구현", "40%", "ZmqRecvThread와 EventDetector 연결 구조 구현. JSON parsing 및 socket 수신 TODO"],
            ["이벤트 감지", "초안 구현", "65%", "목 각도/EAR streak 기반 이벤트 트리거 구조 구현"],
            ["이벤트 클립", "뼈대 구현", "45%", "Claim Check 인터페이스 구현. 실제 JPEG sequence/MP4 인코딩 TODO"],
            ["JSONL 로그", "초안 구현", "60%", "AnalysisResult/Event metadata JSONL 생성 구조 구현. WinHTTP POST TODO"],
            ["알림", "뼈대 구현", "45%", "AlertManager, AlertQueue, AlertDispatchThread 구성. MFC 팝업/Arduino Serial TODO"],
            ["빌드 검증", "부분 확인", "35%", ".vcxproj XML 및 파일 참조 검증 완료. MSBuild/CMake 실제 빌드 미실행"],
        ],
        widths=[2.8, 2.3, 1.8, 8.6],
    )

    add_heading(doc, "6. 현재 코드 기준 주요 파일", 1)
    add_table(
        doc,
        ["폴더", "역할", "대표 파일"],
        [
            ["include/core", "공용 동기화/스레드 기반", "RingBuffer.h, ThreadSafeQueue.h, WorkerThreadPool.h"],
            ["include/model", "데이터 모델", "Frame.h, AnalysisResult.h, PostureEvent.h, Alert.h"],
            ["include/capture", "웹캠 캡처", "CaptureThread.h"],
            ["include/render", "화면 렌더링", "RenderThread.h, OverlayPainter.h"],
            ["include/network", "ZMQ/HTTP JSONL/Claim Check 통신", "IFrameSender.h, ILogSink.h, IEventClipStore.h, ClientTransportConfig.h"],
            ["include/event", "이벤트 감지 및 클립 윈도우", "EventShadowBuffer.h, PostureEventDetector.h"],
            ["include/alert", "팝업/아두이노 알림", "AlertManager.h, AlertDispatchThread.h"],
            ["include/analysis", "분석 엔진 추상화", "IPoseAnalyzer.h, ZmqPoseAnalyzer.h, LocalMediaPipePoseAnalyzer.h"],
        ],
        widths=[3.0, 4.2, 8.3],
    )

    add_heading(doc, "7. 장점", 1)
    for text in [
        "렌더링과 AI 분석이 분리되어 AI 서버가 느려져도 화면 출력은 계속 유지된다.",
        "링버퍼를 사용하므로 지연이 누적되지 않고 최신 프레임 중심으로 복구된다.",
        "AI 서버 전송은 30fps 전체가 아니라 5fps 샘플링이므로 네트워크/인코딩 부하를 줄인다.",
        "이벤트 영상은 상시 전송하지 않고 이벤트 발생 시 Claim Check 방식으로 처리하여 메인서버 트래픽을 억제한다.",
        "통신 책임을 인터페이스로 분리하여 ZMQ, HTTP JSONL, 클립 저장 방식을 설정 기반으로 교체할 수 있다.",
    ]:
        add_body(doc, "• " + text)

    add_heading(doc, "8. 우려되는 병목 및 보완 계획", 1)
    add_table(
        doc,
        ["리스크", "영향", "보완 계획"],
        [
            ["AI 결과 지연", "랜드마크 오버레이가 현재 화면보다 늦게 표시될 수 있음", "Frame/AnalysisResult timestamp 기반 매칭 캐시 구현"],
            ["이벤트 클립 시점 불일치", "AI 응답이 늦으면 이벤트 순간보다 뒤쪽 프레임이 추출될 수 있음", "snapshot_around(timestamp_ms) 방식으로 개선"],
            ["CPU->GPU 업로드 비용", "고해상도에서 렌더 프레임 타임 증가 가능", "D2DFrameUploader에서 ID2D1Bitmap 및 BGRA 변환 버퍼 재사용"],
            ["JPEG 인코딩 비용", "전송 스레드 CPU 점유율 증가 가능", "전송 해상도 축소, JPEG quality 조정, WorkerThreadPool 활용"],
            ["스레드 종료", "blocking wait로 종료 시 deadlock 가능", "ThreadSafeQueue close(), optional 반환 정책 도입"],
            ["업로드 실패", "JSONL/클립 메타데이터 유실 가능", "로컬 spool 파일 저장 후 재전송 정책 추가"],
        ],
        widths=[3.5, 5.3, 6.7],
    )

    add_heading(doc, "9. 다음 작업 계획", 1)
    add_table(
        doc,
        ["우선순위", "작업", "목표"],
        [
            ["1", "Direct2D 렌더링 리소스 구현", "D2DDeviceResources, D2DFrameUploader 추가 및 실제 화면 출력 안정화"],
            ["2", "ZMQ 실제 송수신 구현", "ZmqFrameSender socket 송신, ZmqRecvThread JSON parsing 연결"],
            ["3", "timestamp 기반 분석 결과 매칭", "렌더 프레임과 AI 결과의 시간 불일치 최소화"],
            ["4", "EventShadowBuffer 개선", "snapshot_around(timestamp_ms) 구현"],
            ["5", "HTTP JSONL 업로드 구현", "WinHTTP 기반 application/x-ndjson POST"],
            ["6", "Claim Check 클립 저장 구현", "JPEG sequence 또는 MP4 저장 후 clip_ref 전송"],
            ["7", "알림 출력 구현", "MFC 팝업/toast 및 Arduino serial command 연결"],
            ["8", "CMake/VS 빌드 검증", "OpenCV, ZeroMQ 경로 설정 후 Debug x64 빌드 확인"],
        ],
        widths=[2.0, 5.5, 8.0],
    )

    add_heading(doc, "10. 요약", 1)
    add_body(doc, "현재 client는 최종 기능 완성 단계가 아니라, 실시간 머신비전 앱의 핵심 아키텍처를 먼저 고정한 상태이다. 캡처, 렌더링, AI 전송, 결과 수신, 이벤트 추출, 로그 전송, 알림 처리를 독립 컴포넌트로 나누었고, 통신은 ZMQ + HTTP JSONL + Claim Check 조합으로 설계했다.")
    add_body(doc, "진행률을 전체 기준으로 보면 약 55~60% 수준이다. 구조와 인터페이스는 상당 부분 정리되었고, 다음 단계는 Direct2D 실제 렌더링, ZMQ socket 연결, WinHTTP 업로드, 이벤트 클립 인코딩을 채워 넣는 구현 단계이다.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
